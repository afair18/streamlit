import streamlit as st
import time
from streamlit_option_menu import option_menu
from PIL import Image
import pandas as pd
import numpy as np





# 사이드바 메뉴
with st.sidebar:
    choice = option_menu("Menu", ["세팅","출력", "폼", "차트","MYSQL","기타기능","langchain","모두보기"],
    icons=['gear','view-stacked', 'ui-checks', 'bar-chart','database-check','cpu','robot'],
    menu_icon="app-indicator", default_index=0,
    styles={
        "container": {"padding": "4!important", "background-color": "#fafafa"},
        "icon": {"color": "black", "font-size": "20px"},
        "nav-link": {"font-size": "14px", "text-align": "left", "margin":"0px", "--hover-color": "#fafafa"},
        "nav-link-selected": {"background-color": "#08c7b4"},
    }
    )



# 모듈세팅 및 기본설치
def setting():
    st.header("모듈 설치")
    code = '''
    # streamlit 설치
    pip install streamlit
    import streamlit as st

    # streamlit 실행
    streamlit run sample.py

    # 차트 설치
    import pandas as pd
    import numpy as np

    # 메뉴 왼쪽 icons 참고
    https://icons.getbootstrap.com/

    # MYSQL 설치+사용
    pip install mysql-connector-python
    import mysql.connector


    # lang chain 설치
    pip install langchain

    #  메뉴 설치
    pip install streamlit-option-menu

    '''

    code2 = '''
    # 메뉴 구성 샘플코드
    import streamlit as st
    from streamlit_option_menu import option_menu

    # 사이드바 메뉴
    with st.sidebar:
        choice = option_menu("Menu", ["메뉴1","메뉴2"],
        icons=['gear','view-stacked', 'ui-checks', 'bar-chart','database-check','cpu','robot'],
        menu_icon="app-indicator", default_index=0,
        styles={
            "container": {"padding": "4!important", "background-color": "#fafafa"},
            "icon": {"color": "black", "font-size": "20px"},
            "nav-link": {"font-size": "14px", "text-align": "left", "margin":"0px", "--hover-color": "#fafafa"},
            "nav-link-selected": {"background-color": "#08c7b4"},
        }
        )

    # 각 메뉴에 대한 내용
    def menu1():
        st.write('메뉴1 입니다.')

    def menu2():
        st.write('메뉴2 입니다.')

    # 메뉴연결
    if choice == "메뉴1":
        menu1()
    elif choice == "메뉴2":
        menu2()
    else:
        menu1()
    '''

    st.code(code)
    st.code(code2)

    st.markdown("---")

#기본출력
def view():
    # 출력부분
        
    st.title("텍스트 출력-b1-병합")
    st.header("Header")
    st.subheader("subheader")
    st.write('Hello, *World!* :sunglasses:')
    st.caption('Balloons. Hundreds of them...')
    st.latex(r''' e^{i\pi} + 1 = 0 ''')
    st.write(['st', 'is <', 3]) 
    code = '''
    import streamlit as st
    st.write('Hello, *World!* :sunglasses:')
    '''
    st.code(code)

    st.markdown("---")

    st.header("테이블출력")
    import pandas as pd
    st.write(pd.DataFrame({
        'first column': [1, 2, 3, 4],
        'second column': [10, 20, 30, 40],
    }))

    code = '''
    import pandas as pd

    st.write('테이블출력 예제')
    st.write(pd.DataFrame({
        'first column': [1, 2, 3, 4],
        'second column': [10, 20, 30, 40],
    }))
    '''
    st.code(code)

    st.markdown("---")
    st.header("컬럼 출력")
    col1,col2 = st.columns([2,3])
    # 공간을 2:3 으로 분할하여 col1과 col2라는 이름을 가진 컬럼을 생성합니다.  

    with col1 :
        # column 1 에 담을 내용
        st.title('here is column1')
    with col2 :
        # column 2 에 담을 내용
        st.title('here is column2')
        st.checkbox('this is checkbox1 in col2 ')


    # with 구문 말고 다르게 사용 가능 
    col1.subheader(' i am column1  subheader !! ')
    col2.checkbox('this is checkbox2 in col2 ') 
    #=>위에 with col2: 안의 내용과 같은 기능을합니다


    st.markdown("---")
    st.header("탭 출력")
    # 탭 생성 : 첫번째 탭의 이름은 Tab A 로, Tab B로 표시합니다. 
    tab1, tab2= st.tabs(['Tab A' , 'Tab B'])

    with tab1:
    #tab A 를 누르면 표시될 내용
        st.subheader('hello')
        
    with tab2:
    #tab B를 누르면 표시될 내용 
        st.subheader('hi')

    st.markdown("---")
    st.header("사이드바")

    


    # 사이드바에 체크박스, 버튼등 추가할 수 있습니다! 


    # st.markdown("---")
    # st.header("로딩상태 출력")
    # latest_iteration = st.empty()
    # bar = st.progress(0)
    # for i in range(100):
    #   latest_iteration.text(f'Iteration {i+1}')
    #   bar.progress(i + 1)
    #   time.sleep(0.05)
    # st.balloons()

    # st.markdown("---")
    # st.header("웨이팅 출력")
    # with st.spinner('Wait for it...'):
    #   time.sleep(5)
    #   st.success('Done!')


    st.markdown("---")
    st.header("이미지 출력")
    img = Image.open('sample_img.jpg')
    st.image(img)


#    st.markdown("---")
#    st.header("동영상 출력")
#    video_file = open('sample_movie.mp4', 'rb')
#    st.video(video_file)

    st.markdown("---")
    st.header("오디오 출력")
    audio_file = open('sample_audio.mp3', 'rb')
    st.audio( audio_file.read() , format='audio/mp3')

    st.markdown("---")
    # CSS를 사용해 버튼 스타일 정의
    st.header("css 버튼 출력+파이썬 변수")
    button_style = """
    <style>
    .custom-button {
        display: inline-block;
        padding: 0.5rem 1rem;
        font-size: 1rem;
        color: white;
        background-color: gray;
        border: none;
        border-radius: 4px;
        cursor: pointer;
    }
    .custom-button:hover {
        background-color: black;
    }
    </style>
    """

    # Streamlit 페이지에 CSS 적용
    st.markdown(button_style, unsafe_allow_html=True)


    py_variable = "파이썬 변수"
    # HTML 버튼 생성
    button_html = f'<button onclick="handleClick()" class="custom-button">{py_variable}</button>'
    st.markdown(button_html, unsafe_allow_html=True)

#==================================================================================================

# 폼 출력
def form():
    st.markdown("---")
    st.header("텍스트 입력 폼")
    # 1. 이름 입력 받기
    name = st.text_input('이름을 입력하세요!')

    if name != '' :
        st.subheader(name + '님 안녕하세요??')
    # 2. 입력 글자 갯수 제한
    address = st.text_input('주소를 입력하세요!', max_chars=10)

    # 여러 행을 입력가능하도록
    message = st.text_area('메세지를 입력하세요', height=3)
    st.subheader(message)

    code ='''
    name = st.text_input('이름을 입력하세요!')
    address = st.text_input('주소를 입력하세요!', max_chars=10)
    message = st.text_area('메세지를 입력하세요', height=3)
    '''
    st.code(code)



    st.markdown("---")
    st.header("체크박스 선택")
    # 체크박스 선택
    agree = st.checkbox('동의합니다')
    if agree :
        st.write('동의하셨습니다.')
    else :
        st.write('동의하지 않았습니다.')

    code = '''
    agree = st.checkbox('동의합니다')
    if agree :
        st.write('동의하셨습니다.')
    else :
        st.write('동의하지 않았습니다.')
    '''
    st.code(code)


    # 라디오버튼 선택예제
    st.markdown("---")
    st.header("라디오 버튼")
    

    options = ['Option 1', 'Option 2', 'Option 3']
    selected_option = st.radio('Select an option', options)
    st.write('You selected:', selected_option)


    code = '''
    options = ['Option 1', 'Option 2', 'Option 3']
    selected_option = st.radio('Select an option', options)
    st.write('You selected:', selected_option)
    '''
    st.code(code)

    # 숫자 입력 폼
    st.markdown("---")
    st.header("숫자입력 폼")
    # 4. 숫자 입력, 정수
    st.number_input('숫자 입력', 1, 100)
    # 5. 숫자 입력, 실수
    st.number_input('실수 입력', 1.0, 100.0)


    code = '''
    st.number_input('숫자 입력', 1, 100)
    st.number_input('실수 입력', 1.0, 100.0)
    '''
    st.code(code)

    # 날짜 입력 폼
    st.markdown("---")
    st.header("날짜입력 폼")
    my_date = st.date_input('약속날짜') # 디폴트로 오늘 날짜가 찍혀 있다.
    st.write(my_date)

    # 요일 출력
    st.write( my_date.weekday() )
    st.write( my_date.strftime('%A') )

    code = '''
    my_date = st.date_input('약속날짜') 
    st.write(my_date)
    
    # 요일 출력
    st.write( my_date.weekday() )
    st.write( my_date.strftime('%A') )

    '''
    st.code(code)

    # 시간 입력폼
    st.markdown("---")
    st.header("시간 입력폼")
    my_time = st.time_input('시간 선택')
    st.write(my_time)
    code = '''
    my_time = st.time_input('시간 선택')
    st.write(my_time)
    '''
    st.code(code)



    # 다중선택
    st.markdown("---")
    st.header("다중 선택")
    option_list = ['짜장면', '짬뽕', '탕수육']
    option = st.multiselect('메뉴를 선택하세요', option_list)
    st.write(option)

    code = '''
    option_list = ['짜장면', '짬뽕', '탕수육']
    option = st.multiselect('메뉴를 선택하세요', option_list)
    st.write(option)
    '''
    st.code(code)


    # 슬라이더 선택
    st.markdown("---")
    st.header("슬라이더 선택")
    age = st.slider('나이', 1, 120, 30, 5)
    st.text('제가 선택한 나이는 {} 입니다.'.format(age))
    code = '''
    age = st.slider('나이', 1, 120, 30, 5)
    st.text('제가 선택한 나이는 {} 입니다.'.format(age))
    '''
    st.code(code)



    # 엑셀파일 출력
    st.markdown("---")
    st.header("엑셀파일 출력")
    df = pd.read_csv('sample_csv.csv')
    with st.expander('데이터프레임 보기') :
        st.dataframe(df)
    
    code = '''
    df = pd.read_csv('sample_csv.csv')
    with st.expander('데이터프레임 보기') :
        st.dataframe(df)
    '''
    st.code(code)




    # 색상입력 폼
    st.markdown("---")
    st.header("색상 입력폼")
    color = st.color_picker('색을 선택하세요')
    st.write(color)
    code = '''
    color = st.color_picker('색을 선택하세요')
    st.write(color)
    '''
    st.code(code)


#차트 출력
def chart():
    # 차트 출력
    # 바차트
    chart_data = pd.DataFrame(
    {
        "col1": list(range(20)) * 3,
        "col2": np.random.randn(60),
        "col3": ["A"] * 20 + ["B"] * 20 + ["C"] * 20,
    }
    )

    st.bar_chart(chart_data, x="col1", y="col2", color="col3")


    #영역차트
    chart_data = pd.DataFrame(np.random.randn(20, 3), columns=["a", "b", "c"])
    st.area_chart(chart_data)


    # 라인차트
    chart_data = pd.DataFrame(np.random.randn(20, 3), columns=["a", "b", "c"])
    st.line_chart(chart_data)


    # scatter 차트
    chart_data = pd.DataFrame(np.random.randn(20, 3), columns=["a", "b", "c"])
    st.scatter_chart(chart_data)

    # 지도
    df = pd.DataFrame(
        np.random.randn(1000, 2) / [50, 50] + [37.76, -122.4],
        columns=['lat', 'lon'])
    st.map(df)





def mysql():
    st.header("MYSQL 접속")
    code = '''
    import mysql.connector
    # MySQL 데이터베이스 연결 설정
    db = mysql.connector.connect(
        host="localhost",
        user="your_username",
        password="your_password",
        database="your_database"
    )

    # MySQL 커서 생성
    cursor = db.cursor()
    '''
    st.code(code)
    st.markdown("---")

    st.header("MYSQL 데이터 추가")
    code = '''
    new_data = st.text_input("새로운 데이터 입력:")
    if st.button("추가"):
        insert_query = "INSERT INTO your_table_name (column_name) VALUES (%s)"
        cursor.execute(insert_query, (new_data,))
        db.commit()
        st.success("데이터가 추가되었습니다.")
    '''
    st.code(code)
    st.markdown("---")

    st.header("MYSQL 데이터 조회")
    code = '''
    # 데이터 조회
    cursor.execute("SELECT * FROM your_table_name")
    result = cursor.fetchall()
    for row in result:
        st.write(row)
    '''
    st.code(code)
    st.markdown("---")

    st.header("MYSQL 데이터 수정")
    code = '''
    # 데이터 수정
    update_data = st.text_input("수정할 데이터 입력:")
    if st.button("수정"):
        update_query = "UPDATE your_table_name SET column_name = %s WHERE column_name = %s"
        cursor.execute(update_query, (update_data, update_data))
        db.commit()
        st.success("데이터가 수정되었습니다.")
    '''
    st.code(code)
    st.markdown("---")


def etc():
    st.header("웹페이지 파싱")
    code = '''
    import requests
    from bs4 import BeautifulSoup

    # 네이버 금융에서 종목코드 넣으면 가격정보 가져오기

    codes = ['096530', '010130'] # 종목코드 리스트
    prices = [] # 가격정보가 담길 리스트

    for code in codes:
        url = 'https://finance.naver.com/item/main.nhn?code=' + code
        response = requests.get(url)

        response.raise_for_status()

        html = response.text

        soup = BeautifulSoup(html, 'html.parser')
        today = soup.select_one('#chart_area > div.rate_info > div')
        price = today.select_one('.blind')
        prices.append(price.get_text())
        print(prices)
'''
    st.code(code)

    import requests
    from bs4 import BeautifulSoup

    # 네이버 금융에서 종목코드 넣으면 가격정보 가져오기
    
    codes = ['096530', '010130'] # 종목코드 리스트
    prices = [] # 가격정보가 담길 리스트

    for code in codes:
        url = 'https://finance.naver.com/item/main.nhn?code=' + code
        response = requests.get(url)
        response.raise_for_status()

        html = response.text

        soup = BeautifulSoup(html, 'html.parser')
        today = soup.select_one('#chart_area > div.rate_info > div')
        price = today.select_one('.blind')
        prices.append(price.get_text())
        #print(prices)
    st.write("실행결과")
    st.write(prices)
    

    st.header("네이버 자동로그인")
    code = '''
    import pyperclip
    from selenium import webdriver
    from selenium.webdriver.common.keys import Keys
    from selenium.webdriver.common.by import By
    import time

    # 웹드라이버 열기 (네이버 메인 화면)
    driver = webdriver.Chrome()
    driver.get("https://nid.naver.com/nidlogin.login?mode=form&url=https%3A%2F%2Fwww.naver.com")

    time.sleep(3)   # 3초 시간 지연

    # 로그인 창에 아이디/비밀번호 입력
    loginID = "네이버아이디"
    pyperclip.copy(loginID)
    driver.find_element(By.XPATH, '//*[@id="id"]').send_keys(Keys.CONTROL + 'v') # 붙여넣기


    loginPW = "네이버비번"
    pyperclip.copy(loginPW)
    driver.find_element(By.XPATH, '//*[@id="pw"]').send_keys(Keys.CONTROL + 'v') # 붙여넣기


    time.sleep(1)

    # 로그인 버튼 클릭
    driver.find_element(By.XPATH, '//*[@id="log.login"]').click()
    time.sleep(2)

    #로그인 후 '새로운 환경' 알림에서 '나중에 하기' 클릭
    driver.find_element(By.XPATH, '//*[@id="new.dontsave"]').click()

    while(True):
        pass
    '''
    st.code(code)



    st.header("엑셀,워드 다루기")
    code = '''
    import os
    import openpyxl as op #엑셀 모듈
    from docx import Document #워드 모듈

    from openpyxl.styles.fonts import Font
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Cm, Inches
    from docx.oxml.ns import qn
    from docx.shared import Pt



    # 현재 파일 경로
    path = os.path.dirname(os.path.abspath(__file__))

    wb = op.load_workbook(f"{path}/excel.xlsx") #워크북 객체 생성(파일명 : test.xlsx)
    ws = wb.active #활성화 되어있는 시트 설정


    #Sheet의 Cell 속성 사용하기
    data1 = ws.cell(row=1, column=2).value

    #엑셀 인덱스(Range) 사용하기
    data2 = ws["C1"].value


    #위 결과 출력
    print("cell(1,2) : ", data1)
    print('Range("C1"):', data2)



    # 엑셀의 D4 셀에 C의 합(sum 함수 이용)을 넣고 저장
    ws["D4"].value = "=SUM(A4:C4)"
    font_format = Font(size=15, name='굴림', color = 'FF0000',bold = True,italic = True)
    ws["D4"].font = font_format

    # 새로운 시트 생성 ws = wb.create_sheet("연습")
    wb.save(f"{path}/excel.xlsx") # 엑셀 저장

    # 엑셀의 값을 데이터에 넣기
    data = []
    for row in ws.rows:
        data.append(row[1].value) # B열 엑셀 데이터를 리스트에 추가

    print(data)


    # 워드문서 만들기

    # 불러오기
    doc = Document()

    # 글씨입력
    doc.add_heading('제목 입니다', level=1)
    para = doc.add_paragraph('문단에 들어갈 내용입니다.') #문단 추가


    #이미지삽입
    doc.add_picture(f'{path}/excelimg.jpg', width=Cm(10)) #이미지 삽입



    #표만들기
    table = doc.add_table(rows=2, cols=3) #표 추가
    table.style = doc.styles['Table Grid']
    cell = table.cell(0, 1) #0행 1열
    cell.text = '표에 들어갈 내용입니다.' #표에 내용 추가

    #저장하기
    doc.save(f'{path}/word.docx')
    '''
    st.code(code)

    from streamlit.components.v1 import html

    # HTML + JavaScript 코드
    HTML_CODE = """
    <button id="startRecord">녹음 시작</button>
    <button id="stopRecord" disabled>녹음 정지</button>
    <p id="recognitionResult">인식된 텍스트가 여기에 표시됩니다...</p>

    <script>
    var recognition;
    document.getElementById("startRecord").addEventListener("click", function() {
        if ('webkitSpeechRecognition' in window) {
            recognition = new webkitSpeechRecognition();
            recognition.lang = 'ko-KR'; // 한국어 설정
            recognition.continuous = false; // 단일 결과
            recognition.interimResults = false; // 중간 결과 미표시

            recognition.onstart = function() {
                document.getElementById("recognitionResult").textContent = "인식 중...";
            };

            recognition.onresult = function(event) {
                var transcript = event.results[0][0].transcript;
                document.getElementById("recognitionResult").textContent = transcript;
            };

            recognition.onerror = function(event) {
                document.getElementById("recognitionResult").textContent = "인식 에러: " + event.error;
            };

            recognition.onend = function() {
                document.getElementById("startRecord").disabled = false;
                document.getElementById("stopRecord").disabled = true;
            };

            recognition.start();
            document.getElementById("startRecord").disabled = true;
            document.getElementById("stopRecord").disabled = false;
        } else {
            document.getElementById("recognitionResult").textContent = "이 브라우저에서는 음성 인식이 지원되지 않습니다.";
        }
    });

    document.getElementById("stopRecord").addEventListener("click", function() {
        if (recognition) {
            recognition.stop();
        }
    });
    </script>
    """
    # Streamlit 앱에서 HTML 코드 삽입
    st.title('Streamlit 음성 인식 예제')
    html(HTML_CODE, height=200)

    HTML_CODE = """
    <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/cropperjs/1.5.12/cropper.min.css" />
    <style>
        img {
            max-width: 100%;
            height: 300px;
        }
        .container {
            margin-top: 20px;
        }
    </style>
<input type="file" id="imageInput" accept="image/*">
<div class="container">
    <img id="image" style="display:none;">
</div>
<button id="cropButton">자르기</button>

<script src="https://cdnjs.cloudflare.com/ajax/libs/cropperjs/1.5.12/cropper.min.js"></script>
<script>
    let image = document.getElementById('image');
    let input = document.getElementById('imageInput');
    let cropButton = document.getElementById('cropButton');
    let cropper;

    input.addEventListener('change', function(e) {
        const files = e.target.files;
        if (files && files.length > 0) {
            const fileReader = new FileReader();

            fileReader.onload = function() {
                image.src = this.result;
                image.style.display = 'block';

                if (cropper) {
                    cropper.destroy();
                }

                cropper = new Cropper(image, {
                    aspectRatio: 16 / 9,
                    viewMode: 1,
                });
            };

            fileReader.readAsDataURL(files[0]);
        }
    });

    cropButton.addEventListener('click', function() {
        if (!cropper) {
            return;
        }

        const canvas = cropper.getCroppedCanvas();
        if (!canvas) {
            return;
        }

        canvas.toBlob(function(blob) {
            const url = URL.createObjectURL(blob);
            const a = document.createElement('a');
            a.href = url;
            a.download = 'cropped-image.png';
            document.body.appendChild(a);
            a.click();
            document.body.removeChild(a);
            URL.revokeObjectURL(url);
        }, 'image/png');
    });
    </script>
    
    """
    st.title('이미지 자르기 예제')
    html(HTML_CODE,height=400)

def langchain():
    st.header("Langchain")
    
    code = '''
    from langchain.chat_models import ChatOpenAI
    import openai
    import os


    
    # Langchain을 사용하여 모델 연결
    llm = ChatOpenAI(temperature=0,               # 창의성 (0.0 ~ 2.0) 
                 max_tokens=2048,             # 최대 토큰수
                 model_name='gpt-3.5-turbo',  # 모델명
                )

    # 질의내용
    question = '미국의 수도는 뭐야?'

    # 질의
    st.write((f'[답변]: {llm.predict(question)}'))



    # 사용가능한 모델 목록
    openai.api_key = "키값"
    model_list = sorted([m['id'] for m in openai.Model.list()['data']])
    for m in model_list:
        st.write(m)


    # 템플릿 사용
    from langchain.prompts import PromptTemplate
    from langchain.chains import LLMChain
    template = '{area1} 와 {area2} 의 시차는 몇시간이야?'
    prompt = PromptTemplate(template=template, input_variables=['area1', 'area2'])
    chain = LLMChain(prompt=prompt, llm=llm)
    st.write(chain.run(area1='서울', area2='파리'))
    '''
    st.code(code)


def allshow():
    # version: 1.35.0
    # Description: Streamlit sample code
    # 현재 버전 확인
    #streamlit --version

    # Streamlit 업그레이드
    #pip install --upgrade streamlit

    # 업그레이드 후 버전 확인
    #streamlit --version


    import streamlit as st
    import pandas as pd
    import numpy as np


    # 페이지 세팅 상단에 위치


    st.subheader('페이지세팅 / 글씨출력')
    code = '''
    st.set_page_config(
        page_title="Streamlit-Samples",
        page_icon="🧊",
        layout="wide",
        initial_sidebar_state="expanded",
        menu_items={
            'Get Help': 'https://www.extremelycoolapp.com/help',
            'Report a bug': "https://www.extremelycoolapp.com/bug",
            'About': "# This is a header. This is an *extremely* cool app!"
        }
    )
    '''
    st.code(code)


    code = '''
    # 아이콘이 있는 메뉴
    from streamlit_option_menu import option_menu

    with st.sidebar:
        choice = option_menu("Menu", ["메뉴1","메뉴2"],
        icons=['gear','view-stacked', 'ui-checks', 'bar-chart','database-check','cpu','robot'],
        menu_icon="app-indicator", default_index=0,
        styles={
            "container": {"padding": "4!important", "background-color": "#fafafa"},
            "icon": {"color": "black", "font-size": "20px"},
            "nav-link": {"font-size": "14px", "text-align": "left", "margin":"0px", "--hover-color": "#fafafa"},
            "nav-link-selected": {"background-color": "#08c7b4"},
        }
        )

    # 각 메뉴에 대한 내용
    def menu1():
        st.write('메뉴1 입니다.')

    def menu2():
        st.write('메뉴2 입니다.')

    # 메뉴연결
    if choice == "메뉴1":
        menu1()
    elif choice == "메뉴2":
        menu2()
    else:
        menu1()
    '''

    st.code(code)



    sidebar = st.sidebar
    with sidebar:
        st.write('왼쪽 메뉴부분에 출력')
    code = '''
    sidebar = st.sidebar
    with sidebar:
        st.write('왼쪽 메뉴부분에 출력')
    '''
    st.code(code)

    st.title('This is a title')
    st.header('This is a header')
    st.subheader('This is a subheader')
    st.text('This is a text')
    st.write('Hello, world!')

    code = '''
    st.title('This is a title')
    st.header('This is a header')
    st.subheader('This is a subheader')
    st.text('This is a text')
    st.write('Hello, world!')
    '''

    st.code(code)
    st.divider()

    st.subheader('폼입력')

    name = st.text_input('이름 입력')
    address = st.text_area('주소 입력')

    colum = st.columns(2)
    with colum[0]:
        date1 = st.date_input('날짜 입력')

    with colum[1]:
        time1 = st.time_input('시간 입력')

    st.write(pd.DataFrame({'name': [1,2,3,4], 'date1': [5,6,7,8], 'time': [9,10,11,12]}))

    tab = st.tabs(['tab1', 'tab2', 'tab3'])

    with tab[0]:
        st.write('tab1')

    agree = st.checkbox('동의합니다.')

    if agree:
        st.write('동의하셨습니다.')

    on = st.toggle("Activate feature")

    if on:
        st.write("Feature activated!")

    options = ['짜장', '짬뽕', '탕수육']

    selected = st.selectbox('메뉴를 선택하세요.', options)
    st.write('선택한 메뉴:', selected)

    multi_selected = st.multiselect('메뉴를 선택하세요.', options) 
    st.write('선택한 메뉴:', multi_selected)

    number = st.number_input('숫자를 입력하세요.', min_value=0, max_value=100)
    st.write('입력한 숫자:', number)

    slider = st.slider('숫자를 선택하세요.', min_value=0, max_value=100)
    st.write('선택한 숫자:', slider)

    file = st.file_uploader('파일을 업로드하세요.')

    color = st.color_picker('색상을 선택하세요.',value='#00f900')
    st.write('선택한 색상:', color)

    expander = st.expander('내용보기', expanded=True)
    expander.write('여기에는 자세한 내용이 들어갑니다.')

    st.link_button("Go to gallery", "https://streamlit.io/gallery")

    code ='''
    # 텍스트 박스
    name = st.text_input('이름 입력')

    # 텍스트 에어리어
    address = st.text_area('주소 입력')

    # 컬럼 나누기
    colum = st.columns(2)
    with colum[0]:
        date1 = st.date_input('날짜 입력')

    with colum[1]:
        time1 = st.time_input('시간 입력')

    st.write(pd.DataFrame({'name': [1,2,3,4], 'date1': [5,6,7,8], 'time': [9,10,11,12]}))

    #탭
    tab = st.tabs(['tab1', 'tab2', 'tab3'])

    with tab[0]:
        st.write('tab1')

    #체크박스
    agree = st.checkbox('동의합니다.')

    if agree:
        st.write('동의하셨습니다.')

    #토글
    on = st.toggle("Activate feature")

    if on:
        st.write("Feature activated!")

    options = ['짜장', '짬뽕', '탕수육']

    #단일선택
    selected = st.selectbox('메뉴를 선택하세요.', options)
    st.write('선택한 메뉴:', selected)

    #다중선택
    multi_selected = st.multiselect('메뉴를 선택하세요.', options) 
    st.write('선택한 메뉴:', multi_selected)

    #숫자입력
    number = st.number_input('숫자를 입력하세요.', min_value=0, max_value=100)
    st.write('입력한 숫자:', number)

    # 슬라이드
    slider = st.slider('숫자를 선택하세요.', min_value=0, max_value=100)
    st.write('선택한 숫자:', slider)

    #파일 업로드
    file = st.file_uploader('파일을 업로드하세요.')

    # 컬러선택
    color = st.color_picker('색상을 선택하세요.',value='#00f900')
    st.write('선택한 색상:', color)

    expander = st.expander('내용보기')
    expander.write('여기에는 자세한 내용이 들어갑니다.')
    expander.image('https://www.google.com/images/branding/googlelogo/1x/googlelogo_color_272x92dp.png', width=200)

    # 링크버튼
    st.link_button("Go to gallery", "https://streamlit.io/gallery")
    '''

    st.code(code)

    st.divider()

    # 이미지, 비디오, 오디오
    st.image('https://www.google.com/images/branding/googlelogo/1x/googlelogo_color_272x92dp.png', width=200)
    st.video('https://www.w3schools.com/html/mov_bbb.mp4')
    st.audio('https://www.w3schools.com/html/horse.mp3')

    code = '''
    # 이미지, 비디오, 오디오
    st.image('https://www.google.com/images/branding/googlelogo/1x/googlelogo_color_272x92dp.png', width=200)
    st.video('https://www.w3schools.com/html/mov_bbb.mp4')
    st.audio('https://www.w3schools.com/html/horse.mp3')
    '''

    st.code(code)

    st.divider()

    # 차트 샘플 코드
    st.subheader('차트')


    st.bar_chart(pd.DataFrame({'a': [1,2,3,4]}))

    chart_data = pd.DataFrame(
        np.random.randn(10, 3),
        columns=['a', 'b', 'c'])

    st.line_chart(chart_data)


    st.map(pd.DataFrame({'lat': [37.5665], 'lon': [126.9780]}))

    code = '''
    # 바차트
    st.bar_chart(pd.DataFrame({'a': [1,2,3,4]}))

    chart_data = pd.DataFrame(
        np.random.randn(10, 3),
        columns=['a', 'b', 'c'])

    # 라인차트
    st.line_chart(chart_data)

    # 지도
    st.map(pd.DataFrame({'lat': [37.5665], 'lon': [126.9780]}))
    '''

   

    code = '''
    # altair 사용하기
    import pandas as pd
    import numpy as np
    import altair as alt

    chart_data = pd.DataFrame(np.random.randn(20, 3), columns=["a", "b", "c"])

    c = (
        alt.Chart(chart_data)
        .mark_circle()
        .encode(x="a", y="b", size="c", color="c", tooltip=["a", "b", "c"])
    )

    st.altair_chart(c, use_container_width=True)
    '''

    st.code(code)

   

    code = '''
    # plotly 사용하기
    import plotly.express as px
    import streamlit as st

    st.subheader("Define a custom colorscale")
    df = px.data.iris()
    fig = px.scatter(
        df,
        x="sepal_width",
        y="sepal_length",
        color="sepal_length",
        color_continuous_scale="reds",
    )

    tab1, tab2 = st.tabs(["Streamlit theme (default)", "Plotly native theme"])
    with tab1:
        st.plotly_chart(fig, theme="streamlit", use_container_width=True)
    with tab2:
        st.plotly_chart(fig, theme=None, use_container_width=True)
    '''

    st.code(code)


    st.divider()

    # 팝오버
    with st.popover("Open popover"):
        st.markdown("Hello World 👋")
        name = st.text_input("What's your name?")
    st.write("Your name:", name)

    code = '''
    # 팝오버
    with st.popover("Open popover"):
        st.markdown("Hello World 👋")
        name = st.text_input("What's your name?")
    st.write("Your name:", name)
    '''

    st.code(code)



    # 토스트박스
    if st.button('토스트박스'):
        st.toast('Hooray!', icon='🎉')

    code = '''
    # 토스트박스
    if st.button('토스트박스'):
        st.toast('Hooray!', icon='🎉')
    '''

    st.code(code)


    # 로딩박스
    with st.status("Downloading data..."):
        st.write("Searching for data...")
        time.sleep(1)
        st.write("Found URL.")
        time.sleep(1)
        st.write("Downloading data...")

    st.button("Rerun")

    code = '''
    # 로딩박스
    with st.status("Downloading data..."):
        st.write("Searching for data...")
        time.sleep(2)
        st.write("Found URL.")
        time.sleep(1)
        st.write("Downloading data...")
        time.sleep(1)

    st.button("Rerun")
    '''

    st.code(code)

    # 모달 다이얼로그+변수
    @st.experimental_dialog("Cast your vote")
    def vote(item):
        st.write(f"Why is {item} your favorite?")
        reason = st.text_input("Because...")
        if st.button("Submit"):
            st.session_state.vote = {"item": item, "reason": reason}
            st.rerun()

    if "vote" not in st.session_state:
        st.write("Vote for your favorite")
        if st.button("A"):
            vote("A")
        if st.button("B"):
            vote("B")
    else:
        f"You voted for {st.session_state.vote['item']} because {st.session_state.vote['reason']}"

    code = '''
    # 모달 다이얼로그+변수
    @st.experimental_dialog("Cast your vote")
    def vote(item):
        st.write(f"Why is {item} your favorite?")
        reason = st.text_input("Because...")
        if st.button("Submit"):
            st.session_state.vote = {"item": item, "reason": reason}
            st.rerun()

    if "vote" not in st.session_state:
        st.write("Vote for your favorite")
        if st.button("A"):
            vote("A")
        if st.button("B"):
            vote("B")
    else:
        f"You voted for {st.session_state.vote['item']} because {st.session_state.vote['reason']}"
        '''

    st.code(code)

    # chat message 샘플
    # 채팅 메시지를 저장할 세션 상태 초기화
    if "chat_messages" not in st.session_state:
        st.session_state.chat_messages = []

    # 채팅 메시지 샘플
    messages = st.container()
    if prompt := st.text_input("Prompt"):
        # 새로운 메시지를 세션 상태에 추가
        st.session_state.chat_messages.append(("assistant", f"Echo: {prompt}"))
        st.session_state.chat_messages.append(("user", prompt))

        # 모든 메시지를 렌더링
        for sender, message in st.session_state.chat_messages:
            messages.chat_message(sender).write(message)

    code = '''
    # chat message 샘플
    # 채팅 메시지를 저장할 세션 상태 초기화
    if "chat_messages" not in st.session_state:
        st.session_state.chat_messages = []

    # 채팅 메시지 샘플
    messages = st.container()
    if prompt := st.text_input("Prompt"):
        # 새로운 메시지를 세션 상태에 추가
        st.session_state.chat_messages.append(("assistant", f"Echo: {prompt}"))
        st.session_state.chat_messages.append(("user", prompt))

        # 모든 메시지를 렌더링
        for sender, message in st.session_state.chat_messages:
            messages.chat_message(sender).write(message)
    '''

    st.code(code)

    picture = st.camera_input("Take a picture")


    # html 넣기
    code = """
    <style>
        h5 {
            color: blue;
            font-size: 10px;
            margin: 10px;
        }
    </style>
    """
    st.code(code, language='html')
    st.html(code)
    st.markdown("<h5>Lorem ipsum</h5>", unsafe_allow_html=True)


    # iframe 넣기
    import streamlit.components.v1 as components
    components.iframe("https://example.com", height=200)

    code = """
    # iframe 넣기
    import streamlit.components.v1 as components
    components.iframe("https://example.com", height=400)
    """

    st.code(code)



    #metric
    col1, col2, col3 = st.columns(3)
    col1.metric("Temperature", "70 °F", "1.2 °F")
    col2.metric("Wind", "9 mph", "-8%")
    col3.metric("Humidity", "86%", "4%")

    code = '''
    #metric
    col1, col2, col3 = st.columns(3)
    col1.metric("Temperature", "70 °F", "1.2 °F")
    col2.metric("Wind", "9 mph", "-8%")
    col3.metric("Humidity", "86%", "4%")
    '''

    st.code(code)

    # json 사용하기
    json_data = {
        'foo': 'bar',
        'baz': 'boz',
        'stuff': [
            'stuff 1',
            'stuff 2',
            'stuff 3',
            'stuff 5',
        ],
    }

    st.json(json_data)

    #jdata foo 출력
    st.write(json_data['foo'])

    code = '''
    # json 사용하기
    json_data = {
        'foo': 'bar',
        'baz': 'boz',
        'stuff': [
            'stuff 1',
            'stuff 2',
            'stuff 3',
            'stuff 5',
        ],
    }

    st.json(json_data)

    #jdata foo 출력
    st.write(json_data['foo'])
    '''

    st.code(code)

    # 카메라 사용샘플 코드
    st.camera_input('Take a picture', key='unique_key')


    code = '''
    # 카메라 사용샘플 코드
    st.camera_input('Take a picture', key='unique_key')
    '''

    st.code(code)

    # 세션 상태 사용하기
    st.text_input("Your name", key="name")
    st.session_state.name

    code = '''
    # 세션 상태 사용하기
    st.text_input("Your name", key="name")
    st.session_state.name
    '''

    st.code(code)

    # 타임라인 사용하기
   
    code = '''
    # 타임라인 사용하기
    #pip install streamlit-timeline
    from streamlit_timeline import timeline

    # load data
    with open('example.json', "r") as f:
        data = f.read()

    # render timeline
    timeline(data, height=800)
    '''

    st.code(code)

    # 몽고 DB CRUD 샘플코드
    #pip install pymongo
    # 몽고DB 윈도우 다운로드 https://www.mongodb.com/try/download/community

    st.divider()

    st.subheader('몽고 DB CRUD 샘플코드')


    code = '''
    # 몽고 DB CRUD 샘플코드
    #pip install pymongo
    # 몽고DB 윈도우 다운로드 https://www.mongodb.com/try/download/community
    import pymongo
    from pymongo import MongoClient

    # connect to MongoDB
    client = MongoClient("mongodb://localhost:27017/")
    db = client["mydatabase"]
    col = db["customers"]

    # insert data
    data = {"name": "John", "address": "Highway 37"}
    col.insert_one(data)

    # find data
    result = col.find_one()
    st.write(result)

    # update data
    query = {"address": "Highway 37"}
    new_values = {"$set": {"address": "Canyon 123"}}
    col.update_one(query, new_values)

    # delete data
    query = {"address": "Highway 37"}
    col.delete_one(query)
    '''

    st.code(code)


    # langchain 샘플코드
    #pip install langchain
    #pip install openai==0.27.0
    import openai
    from langchain.chat_models import ChatOpenAI
    from langchain.prompts import PromptTemplate
    from langchain.chains import LLMChain
    import os

    
    
    llm = ChatOpenAI(temperature=0,               # 창의성 (0.0 ~ 2.0) 
                max_tokens=2048,             # 최대 토큰수
                model_name='gpt-4o',  # 모델명
                )
    template = '{area1} 와 {area2} 의 시차는 몇시간이야?'
    prompt = PromptTemplate(template=template, input_variables=['area1', 'area2'])
    chain = LLMChain(prompt=prompt, llm=llm)

    st.write(chain.run(area1='서울', area2='파리'))


    code = '''
    # langchain 샘플코드
    #pip install langchain
    #pip install openai==0.27.0
    import openai
    from langchain.chat_models import ChatOpenAI
    from langchain.prompts import PromptTemplate
    from langchain.chains import LLMChain
    import os

    
    
    llm = ChatOpenAI(temperature=0,               # 창의성 (0.0 ~ 2.0) 
                max_tokens=2048,             # 최대 토큰수
                model_name='gpt-4o',  # 모델명
                )
    template = '{area1} 와 {area2} 의 시차는 몇시간이야?'
    prompt = PromptTemplate(template=template, input_variables=['area1', 'area2'])
    chain = LLMChain(prompt=prompt, llm=llm)


    st.write(chain.run(area1='서울', area2='파리'))

    '''

    st.code(code)


# 메뉴에 따라 내용이 다르게 나옴 
if choice == "출력":
    view()
elif choice == "폼":
    form()
elif choice == "차트":
    chart()
elif choice == "MYSQL":
    mysql()
elif choice == "세팅":
    setting()
elif choice == "기타기능":
    etc()
elif choice == "langchain":
    langchain()
elif choice == "모두보기":
    allshow()
else:
    view()






